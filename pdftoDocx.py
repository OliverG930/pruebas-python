import os
from pdf2docx import Converter
from docx import Document
import re

# pasas pdf a docx
origen = "C:\\Users\\Desktop\\otrosdatos\\PDF"
destino = "C:\\Users\\Desktop\\otrosdatos\\PDF"

# Crea la carpeta de destino sino existe xd
os.makedirs(destino, exist_ok=True)

def convertir_limpiar_reemplazar(origen, destino):
    # los patrones mas comunes para remover
    HEADERS_FOOTERS_PATTERNS = [
        re.compile(r'^\s*\d+\s*de\s*\d+\s*$'),
        re.compile(r'\d{1,2}/\d{1,2}/\d{4}'),
        re.compile(r'\d{1,2}:\d{2}'),
        re.compile(r'file:///'),

        re.compile(r'^\s*$')
    ]

    for archivo in os.listdir(origen):
        if archivo.lower().endswith(".pdf"):
            ruta_pdf = os.path.join(origen, archivo)
            nombre_docx = archivo.replace(".pdf", ".docx")
            ruta_docx = os.path.join(destino, nombre_docx)


            nombre_sin_extension = os.path.splitext(archivo)[0]

            print(f"Processing: {archivo}...")

            try:

                cv = Converter(ruta_pdf)
                cv.convert(ruta_docx, start=0, end=None)
                cv.close()
                print(f"  Converted {archivo} to DOCX.")


                doc = Document(ruta_docx)

                original_paragraphs = [p for p in doc.paragraphs]


                body = doc._body
                for element in body.iter_inner_content():
                    element.getparent().remove(element)


                content_to_reinsert = []
                found_title_part1 = False
                found_title_part2 = False
                found_intro_from_pdf = False
                title_search_done = False

                for i, p in enumerate(original_paragraphs):
                    text_content = p.text.strip()


                    if not title_search_done and text_content.startswith("Documentación"):
                        found_title_part1 = True
                        continue


                    if not title_search_done and found_title_part1 and text_content.startswith("Completa:"):
                        found_title_part2 = True
                        title_search_done = True
                        continue

                    if text_content.startswith("1. Introducción") and not found_intro_from_pdf:
                        found_intro_from_pdf = True
                        continue


                    content_to_reinsert.append(p)


                doc.add_heading(nombre_sin_extension, level=1)
                print(f"  Added main title: '{nombre_sin_extension}'.")

                doc.add_paragraph()


                doc.add_paragraph("1. Introducción")
                print(f"  Added '1. Introducción' after the title.")


                previous_paragraph_was_empty = False
                for p in content_to_reinsert:
                    text_content = p.text.strip()


                    is_undesired_pattern = False
                    for pattern in HEADERS_FOOTERS_PATTERNS:
                        if pattern.search(text_content):
                            is_undesired_pattern = True
                            break

                    if is_undesired_pattern:
                        continue


                    if not text_content:
                        if previous_paragraph_was_empty:
                            continue
                        else:
                            previous_paragraph_was_empty = True
                    else:
                        previous_paragraph_was_empty = False


                    body._element.append(p._element)

                print(f"  Reinserted remaining content from original PDF, filtering unwanted patterns and excessive blank lines.")


                doc.save(ruta_docx)
                print(f"  Saved cleaned and modified {nombre_docx}.")

            except Exception as e:
                print(f"  Error processing {archivo}: {e}")

                continue

convertir_limpiar_reemplazar(origen, destino)
print("\n✔ Process completed. Files converted, cleaned, and modified.")
print("ℹ️ Important Note: PDF to DOCX conversion is complex and may require manual adjustments. This script improves cleaning, but perfect design reproduction is difficult to achieve automatically.")